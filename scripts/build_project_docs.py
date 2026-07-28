from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "editable"
SOURCES = [
    ROOT / "docs" / "PRD.md",
    ROOT / "docs" / "BUSINESS_RULES.md",
    ROOT / "docs" / "UX_SPECIFICATION.md",
    ROOT / "docs" / "DATA_MODEL.md",
    ROOT / "docs" / "API_SPECIFICATION.md",
    ROOT / "docs" / "ARCHITECTURE.md",
    ROOT / "PLAN.md",
]

INK = RGBColor(0x3A, 0x34, 0x3C)
ACCENT = RGBColor(0x8E, 0x7C, 0xB0)
MUTED = RGBColor(0x6E, 0x67, 0x70)
FONT = "Arial"


def set_font(run, size: float, color=INK, bold: bool = False, italic: bool = False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, 9, MUTED)


def configure_page(doc: Document, label: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = f"MUARÉ · {label}"
    set_font(header.runs[0], 9, MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc: Document, title: str, metadata: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("MUARÉ")
    set_font(run, 12, ACCENT, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_font(run, 25, INK, bold=True)

    for line in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        set_font(run, 10.5, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "8E7CB0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    doc.add_page_break()


def extract_contents(lines: list[str]) -> list[tuple[int, str]]:
    items = []
    for line in lines:
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if match:
            items.append((len(match.group(1)), match.group(2)))
    return items


def add_contents(doc: Document, items: list[tuple[int, str]]):
    doc.add_heading("Содержание", level=1)
    for level, title in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * (level - 2))
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        set_font(run, 10.5, INK)
    doc.add_page_break()


def add_inline_runs(paragraph, text: str):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, 9.5, INK)
            run.font.name = "Courier New"
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, 11, INK, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, 11, INK)


def add_body(doc: Document, lines: list[str]):
    in_code = False
    code_buffer: list[str] = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("\n".join(code_buffer))
        set_font(run, 8.5, INK)
        run.font.name = "Courier New"
        code_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if not line or line in {"---"}:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            if level == 1:
                continue
            doc.add_heading(heading.group(2), level=min(level - 1, 3))
            continue
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, re.sub(r"^[-*]\s+", "", line))
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line.replace("  ", " "))
    flush_code()


def build(source: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    title = next((line[2:] for line in lines if line.startswith("# ")), source.stem)
    metadata = [
        line
        for line in lines[1:8]
        if line and not line.startswith("#") and line != "---"
    ][:3]
    doc = Document()
    configure_styles(doc)
    configure_page(doc, source.stem.replace("_", " "))
    add_cover(doc, title, metadata)
    add_contents(doc, extract_contents(lines))
    add_body(doc, lines)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    path = OUTPUT / f"{source.stem}.docx"
    doc.core_properties.title = title
    doc.core_properties.subject = "MUARÉ project documentation"
    doc.core_properties.author = "MUARÉ project"
    doc.save(path)
    return path


if __name__ == "__main__":
    generated = [build(source) for source in SOURCES]
    print("\n".join(str(path) for path in generated))
